#!/usr/bin/env python3
"""
parse_probeseite.py — DOCX → Python-Zwischenformat

Parst Probeseite_Notker.docx und erzeugt ein strukturiertes Zwischenformat
(Python-Dicts / dataclasses) für die TEI-XML-Generierung.

Regelbasiert: Farben, Merged Cells, Zeilentypen, Vers-Zuordnung.
"""

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import docx
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Pipeline-Normalisierungen (ehemals Errata-Layer)
# ---------------------------------------------------------------------------
#
# Textkorrekturen aus Pfeifer-Reviews werden direkt im Parser angewendet,
# nicht als nachgelagerte Schicht. Damit bleibt die Pipeline linear:
# DOCX -> parse -> classify -> TEI -> JSON.
#
# Wenn Pfeifer eine korrigierte DOCX liefert, in der die Stellen bereits
# stimmen, sind die Replacements No-Ops (str.replace findet nichts).
# Quelle aller Korrekturen: Pfeifer-Review 2026-04-15.

PFEIFER_CORRECTIONS: list[tuple[str, str]] = [
    # Cassiodor V1-2 (German)
    ('In vier Teilen ist dieses Psalms', 'In vier Teile ist dieses Psalms'),
    ('als sie Grüde des Zorns', 'als sie Gründe des Zorns'),
    # Remigius V1-2 (German)
    ('mehr über die Welt der Erde ausgebreitet', 'mehr über das Universum der Welt ausgebreitet'),
    # Augustinus V3-5 (German) — Tippfehler
    ('verplfichtet', 'verpflichtet'),
    ('lasst und Arbeit tun, damit uns nicht', 'lasst und Arbeit tun, damit wir nicht'),
    # Cassiodor V7 (German) — Abstand nach [...]
    ('Vater [...]\u201eIch habe dich heute geboren', 'Vater [...] \u201eIch habe dich heute geboren'),
    # Augustinus V7 (German) — Komma nach weil
    ('nur die Gegenwart, weil was auch immer', 'nur die Gegenwart, weil, was auch immer'),
    # Augustinus 2 V7 (German) — Streichung
    ('womit er das ewige Geschlecht', 'womit das ewige Geschlecht'),
    # Augustinus V8-9 (German) — Komma nach [...]
    ('Vorstellung der Menschen [...] sodass', 'Vorstellung der Menschen [...], sodass'),
    # Cassiodor V8-9 (German) — Tippfehler
    ('in der unerlegenden Natur', 'in der unterlegenen Natur'),
    # Cassiodor V12-13 (German) — Komma raus
    ('gerechten Weg, das heißt, vom himmlichen', 'gerechten Weg, das heißt vom himmlichen'),
    # Haupttext V3-5 ahd. — Wortgrenze
    ('Prechen chádensie íro', 'Prechen cháden sie íro'),
    # nhd. V3-5 — Komma raus
    ('geht ihnen der Wille, gleichermaßen', 'geht ihnen der Wille gleichermaßen'),
    # nhd. V3-5 — Komma rein
    ('Der lebt im Himmel wird ver-', 'Der lebt im Himmel, wird ver-'),
    # nhd. V3-5 — das/dass
    ('war, das sie seine Vorherbestimmung', 'war, dass sie seine Vorherbestimmung'),
    # nhd. V7 — Komma rein
    ('Mein Vater sagte zu mir mein Sohn', 'Mein Vater sagte zu mir, mein Sohn'),
    # nhd. V8-9 — Punkt rein
    ('dein Erbe Welches ist das', 'dein Erbe. Welches ist das'),
    # nhd. V8-9 — Punkt zu Komma
    ('mit eisernem Stab . das ist unbeugsame', 'mit eisernem Stab, das ist unbeugsame'),
    # nhd. V8-9 — Komma rein
    ('mit eisernem Stab das heißt mit unbeugsamem', 'mit eisernem Stab, das heißt mit unbeugsamem'),
    # nhd. V10-11 — Tippfehler
    ('unterdrückt den Köper', 'unterdrückt den Körper'),
    # nhd. V12 — Tippfehler
    ('ihr nicht ableitet vom', 'ihr nicht abgleitet vom'),
]


def apply_corrections(text: str) -> str:
    """Wendet die Pfeifer-Korrekturen auf einen aggregierten Text an.

    Idempotent: erneutes Anwenden ist No-Op. Sicher gegen DOCX-Updates,
    in denen die Korrekturen bereits enthalten sind.
    """
    if not text:
        return text
    for old, new in PFEIFER_CORRECTIONS:
        if old in text:
            text = text.replace(old, new)
    return text


def normalize_whitespace_in_text_nodes(tei_text: str) -> str:
    """Mehrfach-Leerzeichen in Element-Inhalten zusammenfassen.

    Wirkt nur auf Text zwischen Tags (>...<), nicht auf Indentation. Behebt
    DOCX-Artefakte wie "nämlich  seinen Namen" → "nämlich seinen Namen".

    Erhalten bleiben:
      - XML-Indentation (whitespace zwischen </X> und <Y>, kein Element-Inhalt)
      - Einzelne Leerzeichen in Wort-Sequenzen
    """
    import re

    def collapse(m):
        before = m.group(1)
        content = m.group(2)
        after = m.group(3)
        # Nur sammeln, wenn Content tatsächliche Wort-Inhalte hat (nicht nur whitespace)
        if not content.strip():
            return m.group(0)
        # Multi-spaces in Wort-Inhalten zu single space (innerhalb des content)
        collapsed = re.sub(r'  +', ' ', content)
        return before + collapsed + after

    # Match: zwischen > und < — Element-Text-Inhalte
    return re.sub(r'(>)([^<]*)(<)', collapse, tei_text)




# ---------------------------------------------------------------------------
# Dataclasses (Zwischenformat)
# ---------------------------------------------------------------------------

@dataclass
class Run:
    """Ein DOCX-Run mit Text, Farbe und Formatierung."""
    text: str
    color: Optional[str]  # '806000', '00B050', or None (=black)
    bold: bool = False
    italic: bool = False


@dataclass
class Segment:
    """Ein zusammenhängender Textabschnitt gleicher Farbe/Funktion."""
    text: str
    function: str       # 'psalm', 'translation', 'commentary'
    color: Optional[str]
    runs: list = field(default_factory=list)  # Original-Runs für Feinanalyse


@dataclass
class TextLine:
    """Eine Haupttext-Zeile der Probeseite."""
    segments: list       # List[Segment]
    nhd: str = ''        # nhd. Übersetzung (kursiv-Spalte)
    sigles: str = ''     # Siglen-Spalte (z.B. 'G, R')
    line_number: int = 0


@dataclass
class GlossLine:
    """Eine Interlinearglosse."""
    text: str            # ahd./lat. Glossentext
    nhd: str = ''        # nhd. Übersetzung
    sigles: str = ''
    line_number: int = 0


@dataclass
class SourceEntry:
    """Ein Eintrag im Quellenapparat."""
    sigle: str           # 'A', 'C', 'R', 'Br', 'RII', 'N'
    latin: str = ''
    german: str = ''
    bold_spans: list = field(default_factory=list)  # Positionen fetter Wörter


@dataclass
class PsalterWitness:
    """Ein Textzeuge im synoptischen Psaltervergleich."""
    sigle: str           # 'G', 'R', 'H', 'A', 'C'
    name: str = ''       # 'Gallicanum', etc.
    manuscript: str = ''
    text: str = ''


@dataclass
class VerseGroup:
    """Eine Versgruppe (z.B. Verse 1-2) mit allen zugehörigen Daten."""
    verses: str          # Versreferenz, z.B. '1-2'
    verse_numbers: list  # [1, 2]
    lines: list = field(default_factory=list)      # List[TextLine | GlossLine]
    sources: list = field(default_factory=list)     # List[SourceEntry]


@dataclass
class PsalmData:
    """Gesamtdaten für einen Psalm."""
    psalm_number: int = 2
    verse_groups: list = field(default_factory=list)   # List[VerseGroup]
    psalter_witnesses: list = field(default_factory=list)  # List[PsalterWitness]
    wiener_notker: str = ''


# ---------------------------------------------------------------------------
# Farb-Mapping
# ---------------------------------------------------------------------------

COLOR_MAP = {
    '806000': 'psalm',       # olive → Psalmzitation
    '00B050': 'translation',  # grün → ahd. Übersetzung
    None: 'commentary',       # schwarz → Kommentar
}

# In the DOCX, "R II" appears as two runs "R " + "II"
SIGLE_PATTERN = re.compile(r'^(A|C|R|Br|RII|R\s*II|N)$')


# ---------------------------------------------------------------------------
# DOCX-Hilfsfunktionen
# ---------------------------------------------------------------------------

def get_run_color(run) -> Optional[str]:
    """Extrahiert die Schriftfarbe eines Runs."""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        color_el = rpr.find(qn('w:color'))
        if color_el is not None:
            val = color_el.get(qn('w:val'))
            if val and val.upper() != '000000':
                return val
    return None


def get_cell_runs(cell) -> list[Run]:
    """Extrahiert alle Runs einer Zelle mit Farbe und Formatierung."""
    runs = []
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text:  # auch Leerzeichen bewahren
                runs.append(Run(
                    text=run.text,
                    color=get_run_color(run),
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                ))
    return runs


def cell_text(cell) -> str:
    """Gesamttext einer Zelle."""
    return ''.join(p.text for p in cell.paragraphs).strip()


def is_merged_row(row) -> bool:
    """Prüft ob Cols 0-2 identisch sind (merged cells)."""
    cells = row.cells
    if len(cells) < 3:
        return False
    t0 = cell_text(cells[0])
    t1 = cell_text(cells[1])
    t2 = cell_text(cells[2])
    # Alle drei gleich und nicht-leer
    return t0 == t1 == t2 and t0 != ''


def is_empty_row(row) -> bool:
    """Prüft ob alle Zellen leer sind."""
    return all(cell_text(c) == '' for c in row.cells)


# ---------------------------------------------------------------------------
# Runs → Segmente (Farb-basiert zusammenführen)
# ---------------------------------------------------------------------------

def runs_to_segments(runs: list[Run]) -> list[Segment]:
    """Fasst aufeinanderfolgende Runs gleicher Farbe zu Segmenten zusammen."""
    if not runs:
        return []

    segments = []
    current_color = runs[0].color
    current_runs = [runs[0]]
    current_texts = [runs[0].text]

    for run in runs[1:]:
        if run.color == current_color:
            current_runs.append(run)
            current_texts.append(run.text)
        else:
            fn = COLOR_MAP.get(current_color, 'commentary')
            segments.append(Segment(
                text=''.join(current_texts),
                function=fn,
                color=current_color,
                runs=list(current_runs),
            ))
            current_color = run.color
            current_runs = [run]
            current_texts = [run.text]

    # Letztes Segment
    fn = COLOR_MAP.get(current_color, 'commentary')
    segments.append(Segment(
        text=''.join(current_texts),
        function=fn,
        color=current_color,
        runs=list(current_runs),
    ))

    return segments


# ---------------------------------------------------------------------------
# Zeilentyp-Erkennung
# ---------------------------------------------------------------------------

def detect_source_row(row) -> Optional[SourceEntry]:
    """Erkennt eine Quellenapparat-Zeile: Col 0 = einzelne Sigle."""
    cells = row.cells
    col0_text = cell_text(cells[0]).strip()

    # Normalisierung: "R II" → "RII"
    normalized = col0_text.replace(' ', '')

    if not SIGLE_PATTERN.match(normalized):
        return None

    # Sicherheitscheck: Col 1 muss substantiellen Text haben
    if len(cells) < 2:
        return None
    col1_text = cell_text(cells[1]).strip()
    if len(col1_text) < 5:
        return None

    # Latin-Text: Col 1 (Runs ohne Formatierung, aber Bold bewahren)
    latin_text = col1_text

    # Bold-Spans in Col 1 extrahieren
    bold_spans = []
    col1_runs = get_cell_runs(cells[1])
    pos = 0
    for r in col1_runs:
        if r.bold and r.text.strip():
            bold_spans.append((pos, pos + len(r.text), r.text))
        pos += len(r.text)

    # German translation: Col 2+ (kursiv)
    german = ''
    for ci in range(2, len(cells)):
        ct = cell_text(cells[ci]).strip()
        if ct and ct != col1_text:  # Nicht dupliziert
            german = ct
            break

    return SourceEntry(
        sigle=normalized,
        latin=latin_text,
        german=german,
        bold_spans=bold_spans,
    )


def detect_gloss_line(text_runs: list[Run], nhd: str, sigles: str) -> bool:
    """
    Heuristik: Ist diese Zeile eine Interlinearglosse?
    Kriterien:
    - Kurzer Text (≤ 5 Wörter)
    - Nur schwarze Runs (keine olive/grüne Farbe)
    - nhd-Spalte hat ebenfalls kurzen Text
    - Kein Satzzeichen am Ende (Glossen enden nicht mit '.')
    - Keine Klammern (Bibliographische Angaben)
    - Keine Verben im Vollsatzformat
    """
    if not text_runs:
        return False

    full_text = ''.join(r.text for r in text_runs).strip()
    words = full_text.split()

    if len(words) > 5:
        return False

    # Keine farbigen Runs (olive oder grün → keine Glosse)
    if any(r.color in ('806000', '00B050') for r in text_runs):
        return False

    # nhd-Spalte sollte auch kurz sein (Glosse hat kurze Übersetzung)
    nhd_words = nhd.strip().split() if nhd else []
    if len(nhd_words) > 5:
        return False

    # Leerer Text → keine Glosse
    if len(full_text) < 2:
        return False

    # Klammern → bibliographische Angabe, nicht Glosse
    if '(' in full_text or ')' in full_text:
        return False

    # Text mit Punkt am Ende → eher Kommentar als Glosse
    # (Ausnahme: ".i." = id est, das ist eine Glosse)
    if full_text.endswith('.') and not full_text.startswith('.i.'):
        return False

    # Editorischer Auslassungsmarker "[...]" signalisiert Fortsetzung von Fließtext
    # (Kommentar- oder Haupttext), nicht eine kurze Interlinearglosse.
    # Pfeifer 2026-04-15 Review (V6 "ze_gótes sélbes ána-sihte. [...]"): als Haupttext,
    # nicht als Glosse zu klassifizieren.
    if '[...]' in full_text:
        return False

    # Vollständige Sätze mit Verb → Kommentar
    # Heuristik: enthält ein konjugiertes Verb-Muster (ahd. endet oft auf -t, -n)
    # Aber viele Glossen enthalten auch "-" (Silbentrennung) → das ist OK
    # Strenger: wenn nhd leer, dann eher keine nützliche Glosse
    # (außer es ist ein Verweis wie "kerich" → "Gericht?")

    return True


# ---------------------------------------------------------------------------
# Vers-Referenz parsen
# ---------------------------------------------------------------------------

def parse_verse_ref(text: str) -> Optional[tuple[str, list[int]]]:
    """
    Parst Vers-Referenzen wie '2,1-2', '2,3-2,5', '2,6'.
    Text kommt oft verdreifacht aus Merged Cells (z.B. '2,1-22,1-22,1-2').
    Gibt (ref_string, [verse_numbers]) zurück.
    """
    text = text.strip()
    if not text:
        return None

    # Brute-Force: Probiere alle gültigen Vers-Referenzen für Psalm 2 (Verse 1-13)
    # und prüfe ob sie als Substring vorkommen mit Grenzprüfung (kein Digit danach).
    # Das umgeht das Problem mit zusammengeklebten Merged-Cell-Texten.
    candidates = []

    def has_boundary_match(pattern: str) -> bool:
        """Prüft ob pattern als Substring mit nicht-Digit-Grenze vorkommt."""
        start = 0
        while True:
            idx = text.find(pattern, start)
            if idx == -1:
                return False
            end_idx = idx + len(pattern)
            # Nach dem Match darf kein Digit folgen
            if end_idx >= len(text) or not text[end_idx].isdigit():
                return True
            start = idx + 1

    # Ranges (Psalm,Start-Psalm,End)
    for s in range(1, 14):
        for e in range(s + 1, 14):
            pattern = f'2,{s}-2,{e}'
            if has_boundary_match(pattern):
                candidates.append((f'{s}-{e}', list(range(s, e + 1))))

    # Ranges (Psalm,Start-End)
    for s in range(1, 14):
        for e in range(s + 1, 14):
            pattern = f'2,{s}-{e}'
            if has_boundary_match(pattern):
                candidates.append((f'{s}-{e}', list(range(s, e + 1))))

    # Einzelverse (Psalm,Vers)
    for v in range(1, 14):
        pattern = f'2,{v}'
        if has_boundary_match(pattern):
            candidates.append((str(v), [v]))

    if not candidates:
        return None

    # Bevorzuge Ranges über Einzelverse (spezifischste = längste Versliste)
    # Bei Gleichstand: kürzere Textform bevorzugen
    candidates.sort(key=lambda c: (-len(c[1]), len(c[0])))
    return candidates[0]


# ---------------------------------------------------------------------------
# Hauptparser
# ---------------------------------------------------------------------------

def parse_probeseite(docx_path: str) -> PsalmData:
    """Parst die Probeseite und gibt das strukturierte Zwischenformat zurück."""

    doc = docx.Document(docx_path)
    psalm = PsalmData()

    # Phase 1: Dokument-Elemente in Reihenfolge durchgehen
    # und Paragraphen (Vers-Überschriften) von Tabellen trennen
    elements = []
    body = doc.element.body
    table_index = 0

    para_index = 0
    all_paragraphs = doc.paragraphs

    for elem in body:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            # Text über python-docx Paragraph extrahieren (sauberer als XML-Iteration)
            if para_index < len(all_paragraphs):
                para_text = all_paragraphs[para_index].text.strip()
                if para_text:
                    elements.append(('paragraph', para_text))
            para_index += 1
        elif tag == 'tbl':
            if table_index < len(doc.tables):
                elements.append(('table', doc.tables[table_index]))
                table_index += 1

    # Phase 2: Vers-Gruppen bilden
    current_group = None
    groups = []

    for etype, edata in elements:
        if etype == 'paragraph':
            ref = parse_verse_ref(edata)
            if ref:
                ref_str, verse_nums = ref
                current_group = VerseGroup(
                    verses=ref_str,
                    verse_numbers=verse_nums,
                )
                groups.append(current_group)

        elif etype == 'table' and current_group is not None:
            parse_table(edata, current_group, psalm)

    psalm.verse_groups = groups

    # Zeilennummern vergeben
    for group in psalm.verse_groups:
        for i, line in enumerate(group.lines):
            line.line_number = i + 1

    return psalm


def parse_table(table, group: VerseGroup, psalm: PsalmData):
    """Parst eine einzelne Tabelle und fügt Daten zur VerseGroup/PsalmData hinzu."""

    rows = list(table.rows)
    if not rows:
        return

    ncols = max(len(r.cells) for r in rows)

    # Sonderfall: Wiener Notker (1 Spalte, Header enthält "Wiener")
    if ncols <= 2:
        first_cell_text = cell_text(rows[0].cells[0])
        if 'Wiener' in first_cell_text or 'ÖNB' in first_cell_text:
            if len(rows) > 1:
                psalm.wiener_notker = cell_text(rows[1].cells[0])
            return

        # Sonderfall: Psalter-Vergleich (2 Spalten, Header mit Cod.)
        if 'Cod.' in first_cell_text or 'Psalter' in first_cell_text.lower():
            parse_psalter_table(table, psalm)
            return

    for row in rows:
        if is_empty_row(row):
            continue

        cells = row.cells

        # Prüfe auf Quellenapparat-Zeile
        source = detect_source_row(row)
        if source:
            group.sources.append(source)
            continue

        # Prüfe auf Psaltervergleich-Zeilen (Siglen-Header mit G, R, H)
        if is_psalter_header_row(row, ncols):
            parse_psalter_inline(table, rows, row, psalm)
            return  # Rest der Tabelle ist Psaltervergleich

        # Haupttext-Zeilen (Merged Cells)
        if is_merged_row(row):
            parse_haupttext_row(row, group, ncols)
        elif ncols >= 3:
            # Nicht-gemergte Zeile mit ≥3 Spalten: könnte auch Haupttext sein
            # wenn Col 0 substantiellen Text hat
            col0 = cell_text(cells[0])
            if len(col0) > 3 and not SIGLE_PATTERN.match(col0.strip().replace(' ', '')):
                parse_haupttext_row(row, group, ncols)


def parse_haupttext_row(row, group: VerseGroup, ncols: int):
    """Parst eine Haupttext-Zeile (gemergete oder nicht-gemergte)."""
    cells = row.cells

    # Haupttext: Col 0 (bei merged: Col 0 = Col 1 = Col 2)
    text_runs = get_cell_runs(cells[0])
    if not text_runs:
        return

    segments = runs_to_segments(text_runs)

    # nhd-Spalte und Siglen-Spalte ermitteln.
    # Spalten-Layout hängt von der Tabellenbreite und dem Merge-Pattern ab.
    # Merge-Varianten:
    #   5-6 Spalten: Col 0-2 merged, Col 3 = nhd, Col 4 = Siglen
    #   4 Spalten:   Col 0-1 merged, Col 2 = nhd, Col 3 = Siglen
    #   3 Spalten:   Col 0 = Haupttext, Col 1 = nhd, Col 2 = Siglen (wenn nicht merged)
    nhd = ''
    sigles = ''
    haupttext = cell_text(cells[0])

    if ncols >= 5:
        nhd_candidate = cell_text(cells[3])
        sigles_candidate = cell_text(cells[4])
        if nhd_candidate != haupttext:
            nhd = nhd_candidate
        if len(sigles_candidate) <= 15:
            sigles = sigles_candidate
    elif ncols == 4:
        # 4-Spalten: Col 0-1 könnten gemergt sein (Col 0 == Col 1)
        # Dann: Col 2 = nhd, Col 3 = Siglen
        # Oder: Col 0-2 gemergt, Col 3 = nhd (unwahrscheinlich bei ncols=4)
        col0 = cell_text(cells[0])
        col1 = cell_text(cells[1])
        if col0 == col1:
            # Cols 0-1 gemergt → Col 2 = nhd, Col 3 = Siglen
            nhd_candidate = cell_text(cells[2])
            sigles_candidate = cell_text(cells[3])
            if nhd_candidate != haupttext:
                nhd = nhd_candidate
            if len(sigles_candidate) <= 15:
                sigles = sigles_candidate
        else:
            # Nicht gemergt → Col 3 = nhd (Fallback)
            nhd_candidate = cell_text(cells[3])
            if nhd_candidate != haupttext:
                nhd = nhd_candidate
    elif ncols == 3 and not is_merged_row(row):
        nhd = cell_text(cells[1])
        sigles = cell_text(cells[2])

    # Glossen-Erkennung
    if detect_gloss_line(text_runs, nhd, sigles):
        full_text = ''.join(r.text for r in text_runs).strip()
        # Whitespace normalisieren (DOCX merged cells erzeugen Tabs/Mehrfach-Leerzeichen)
        full_text = ' '.join(full_text.split())
        nhd_clean = ' '.join(nhd.split()) if nhd else nhd
        group.lines.append(GlossLine(
            text=full_text,
            nhd=nhd_clean,
            sigles=sigles,
        ))
    else:
        group.lines.append(TextLine(
            segments=segments,
            nhd=nhd,
            sigles=sigles,
        ))


def is_psalter_header_row(row, ncols: int) -> bool:
    """Erkennt die Sigle-Header-Zeile des Psaltervergleichs."""
    if ncols < 3:
        return False
    cells = row.cells
    texts = [cell_text(c).strip() for c in cells[:min(5, ncols)]]
    # Psaltervergleich: Zeile mit nur Siglen (G, R, H)
    sigles_found = sum(1 for t in texts if t in ('G', 'R', 'H', 'A', 'C'))
    return sigles_found >= 3


PSALTER_NAMES = {
    'G': ('Gallicanum', ''),
    'R': ('Romanum', ''),
    'H': ('Hebraicum (iuxta Hebraeos)', 'Bamberg Ms. 44'),
    'A': ('Augustinus-Psalter', 'St. Gallen Cod. 162'),
    'C': ('Cassiodor-Psalter', 'St. Gallen Cod. 200'),
}


def _add_psalter_witness(psalm: PsalmData, sigle: str, text: str = ''):
    """Fügt einen Psalter-Zeugen hinzu, dedupliziert nach Sigle."""
    # Existierenden Zeugen aktualisieren statt duplizieren
    for wit in psalm.psalter_witnesses:
        if wit.sigle == sigle:
            if text and not wit.text:
                wit.text = text
            return
    name, ms = PSALTER_NAMES.get(sigle, (sigle, ''))
    psalm.psalter_witnesses.append(PsalterWitness(
        sigle=sigle, name=name, manuscript=ms, text=text,
    ))


def parse_psalter_inline(table, rows, header_row, psalm: PsalmData):
    """Parst den Psaltervergleich aus einer Tabelle mit Siglen-Header."""
    cells = header_row.cells
    ncols = len(cells)
    sigles = [cell_text(c).strip() for c in cells[:ncols]]

    header_idx = rows.index(header_row)
    for row in rows[header_idx + 1:]:
        if is_empty_row(row):
            continue
        for ci in range(min(ncols, len(row.cells))):
            if ci < len(sigles) and sigles[ci] in PSALTER_NAMES:
                text = cell_text(row.cells[ci]).strip()
                if text and len(text) > 10:
                    _add_psalter_witness(psalm, sigles[ci], text)
        break


def parse_psalter_table(table, psalm: PsalmData):
    """Parst eine eigenständige Psalter-Vergleichstabelle (T11: A-Psalter, C-Psalter)."""
    # Zeile 0: Header mit Handschriften-Info
    if table.rows:
        for cell in table.rows[0].cells:
            ct = cell_text(cell).strip()
            if 'Cod. 162' in ct:
                _add_psalter_witness(psalm, 'A')
            elif 'Cod. 200' in ct:
                _add_psalter_witness(psalm, 'C')

    # Zeile 1: Texte (eine Zelle pro Zeuge, Position = Spalte)
    if len(table.rows) > 1:
        text_row = table.rows[1]
        # Zuordnung: Spalte 0 → erster Header-Zeuge, Spalte 1 → zweiter
        header_sigles = []
        for cell in table.rows[0].cells:
            ct = cell_text(cell).strip()
            if 'Cod. 162' in ct:
                header_sigles.append('A')
            elif 'Cod. 200' in ct:
                header_sigles.append('C')

        for ci, sigle in enumerate(header_sigles):
            if ci < len(text_row.cells):
                ct = cell_text(text_row.cells[ci]).strip()
                if ct and len(ct) > 20:
                    _add_psalter_witness(psalm, sigle, ct)


# ---------------------------------------------------------------------------
# Ausgabe / Debug
# ---------------------------------------------------------------------------

def print_summary(psalm: PsalmData):
    """Gibt eine Zusammenfassung des geparsten Psalms aus."""
    print(f'=== Psalm {psalm.psalm_number} ===')
    print(f'Versgruppen: {len(psalm.verse_groups)}')

    total_text = 0
    total_gloss = 0
    total_sources = 0

    for vg in psalm.verse_groups:
        text_lines = [l for l in vg.lines if isinstance(l, TextLine)]
        gloss_lines = [l for l in vg.lines if isinstance(l, GlossLine)]
        total_text += len(text_lines)
        total_gloss += len(gloss_lines)
        total_sources += len(vg.sources)

        print(f'\n--- Verse {vg.verses} (V. {vg.verse_numbers}) ---')
        print(f'  Textzeilen: {len(text_lines)}, Glossen: {len(gloss_lines)}, '
              f'Quellen: {len(vg.sources)}')

        for line in vg.lines:
            if isinstance(line, TextLine):
                seg_summary = ' | '.join(
                    f'[{s.function}] {s.text[:50]}' for s in line.segments
                )
                print(f'  L{line.line_number}: {seg_summary}')
                if line.nhd:
                    print(f'    nhd: {line.nhd[:60]}')
                if line.sigles:
                    print(f'    sigles: {line.sigles}')
            elif isinstance(line, GlossLine):
                print(f'  G{line.line_number}: GLOSSE "{line.text}" → "{line.nhd}"')

        for src in vg.sources:
            print(f'  SRC [{src.sigle}]: {src.latin[:60]}...')
            if src.german:
                print(f'    de: {src.german[:60]}...')

    print(f'\n=== Summe: {total_text} Textzeilen, {total_gloss} Glossen, '
          f'{total_sources} Quelleneinträge ===')

    print(f'\nPsalter-Zeugen: {len(psalm.psalter_witnesses)}')
    for wit in psalm.psalter_witnesses:
        print(f'  [{wit.sigle}] {wit.name or "(Name fehlt)"}: '
              f'{wit.text[:60] if wit.text else "(Text fehlt)"}...')

    print(f'\nWiener Notker: {"Ja" if psalm.wiener_notker else "Nein"} '
          f'({len(psalm.wiener_notker)} Zeichen)')


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    docx_path = Path(__file__).parent.parent / 'data' / 'Probeseite_Notker.docx'
    if not docx_path.exists():
        print(f'FEHLER: {docx_path} nicht gefunden')
        sys.exit(1)

    psalm = parse_probeseite(str(docx_path))
    print_summary(psalm)
    return psalm


if __name__ == '__main__':
    main()
